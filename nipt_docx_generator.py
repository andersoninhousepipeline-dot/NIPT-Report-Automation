"""
DOCX Report Generator for NIPS (NIPT) Reports - Professional Suite
Version: 3.0 (100% Accuracy Refinement)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys
import base64
import tempfile
from io import BytesIO
from datetime import datetime
from nipt_assets import HEADER_LOGO_B64, FOOTER_BANNER_B64

class NIPTDocxGenerator:
    """Generates DOCX reports for NIPS with 100% PDF Parity"""
    
    COLORS = {
        'blue_header': '1F497D',
        'low_risk': '15803D',
        'high_risk': 'B91C1C',
        'grey_text': '666666',
        'patient_info_bg': 'F1F1F7',
        'results_header_bg': 'F9BE8F',
        'grey_bg': 'F2F2F2'
    }

    # Static Content
    METHODOLOGY_TEXT = """Maternal whole blood sample was taken from the pregnant mother after 10-week gestation with no risk to the fetus. The circulating cell-free placental DNA was isolated and purified, which was then converted into genomic DNA library using Yourgene cfDNA Library prep kit. This was followed by automated size selection of fragment lengths by QS250 for enriching fetal fraction. The enriched sample pool was subjected to high throughput sequencing using Ion GeneStudio S5 Plus™ System. Finally, analysis was performed using Sage Link V2."""
    
    LIMITATIONS_TEXT = [
        "Test performance is valid only for full chromosomal aneuploidies involving all the autosomes and sex chromosomes with an accuracy of up to 99% for aneuploidies pertaining to chromosomes 13, 18 and 21. The method is intended for use in pregnant women who are at least 10+ weeks pregnant. The method is suitable for both singleton and twin pregnancies. The accuracy may be slightly lower in twin pregnancies due to multiple sources of fetal DNA. Patients with malignancy or a history of malignancy, patients with bone marrow or organ transplant, patients pregnant with more than 2 fetuses are not eligible for the test.",
        "<b>For vanishing twins, sampling must be performed 8 weeks after vanishing event.</b> The test is not intended and not validated for mosaicism, triploidy, partial trisomy/monosomy, uniparental disomy or translocations. Other genomic abnormalities not covered by this assay may impact the phenotype.",
        "A high risk result for twin pregnancies indicates high probability for the presence of at least one affected fetus.",
        "<b>Pregnant women with a low risk test result do not ensure an unaffected pregnancy as this is a screening test.</b> Although this test is accurate, there is still a small possibility for false positive or false negative results. This may be caused by technical and/or biological limitations, including but not limited to confined placental mosaicism (CPM) or other types of mosaicism, maternal constitutional or somatic chromosomal abnormalities, residual cfDNA from a vanished twin or other rare molecular events.",
        "<b>This test is not a diagnostic but a screening test and results should be considered in the context of other clinical criteria.</b> Clinical correlation with ultrasound findings, and other clinical data and tests is recommended.",
        "<b>If definitive diagnosis is desired, invasive prenatal testing by chorionic villus biopsy (CVS) or amniocentesis is necessary.</b> The referral clinician is responsible for counselling before and after the test including the provision of advice regarding the need for additional invasive genetic testing."
    ]
    
    SIGNATURES = [
        {"name": "Dr. Sachin D. Honguntikar", "title": "Head – Anderson Clinical Genetics"},
        {"name": "Dr. Suriyakumar.G", "title": "Director"}
    ]

    def __init__(self, output_path, with_logo=True):
        self.output_path = output_path
        self.doc = Document()
        self._setup_page(with_logo)

    def _setup_page(self, with_logo=True):
        section = self.doc.sections[0]
        section.page_width = Pt(612)
        section.page_height = Pt(792)
        section.top_margin = Pt(40)
        section.bottom_margin = Pt(50)
        section.left_margin = Pt(50)
        section.right_margin = Pt(50)
        
        # Add Footer with Page Number
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run("Page ")
        run.font.name = 'Gill Sans MT'
        run.font.size = Pt(11)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        p._p.append(fldChar1)
        p._p.append(instrText)
        p._p.append(fldChar2)
        p._p.append(fldChar3)
        
        run2 = p.add_run(" of 6\n\n\n")
        run2.font.name = 'Gill Sans MT'
        run2.font.size = Pt(11)
        
        if with_logo:
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                    tmp_file.write(base64.b64decode(FOOTER_BANNER_B64))
                    footer_path = tmp_file.name
                run_img = p.add_run()
                run_img.add_picture(footer_path, width=Inches(6.5))
                os.unlink(footer_path)
            except: pass
        else:
            # No address/contact text required per user request
            pass

    def _add_logo(self):
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                logo_data = base64.b64decode(HEADER_LOGO_B64)
                tmp_file.write(logo_data)
                logo_path = tmp_file.name
            
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(logo_path, width=Inches(6.5), height=Inches(1.15))
            os.unlink(logo_path)
        except: pass

    def _add_header_text(self, page_num):
        if page_num == 1:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p.add_run("Non-Invasive Prenatal Screening (NIPS)")
            run2.font.name = 'Gill Sans MT Bold'
            run2.font.size = Pt(17)
            run2.font.color.rgb = RGBColor.from_string(self.COLORS['blue_header'])

    def generate(self, data, z_scores, with_logo=True):
        # Page 1
        if with_logo: self._add_logo()
        self._add_header_text(1)
        
        # Patient Info Grid MUST be first
        self._add_patient_grid(data)
        
        # PNDT disclaimer
        pndt = self.doc.add_paragraph()
        pndt_run = pndt.add_run("This test does not reveal sex of the fetus & confers to PNDT act, 1994")
        pndt_run.italic = True
        pndt_run.bold = True
        pndt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:fill'), self.COLORS['grey_bg'])
        pndt._p.get_or_add_pPr().append(shading_elm)
        
        self._add_section_header("Test Indication")
        self.doc.add_paragraph(data.get('indication', 'To screen for chromosomal aneuploidies'))
        
        self._add_section_header("Test Performed")
        self.doc.add_paragraph("NIPS for 23 pairs of chromosomes (including sex chromosomal aneuploidies)")
        
        ff = z_scores.get('fetal_fraction', 0)
        overall_risk = "Low risk"
        is_high = any(abs(z_scores.get(f'chr{i}', 0)) > (2.8 if i in [13,18,21] else 6.0) for i in range(1, 23))
        if is_high or abs(z_scores.get('chrX', 0)) > 6.0: overall_risk = "High risk"
        
        self._add_section_header("Result Summary")
        res_table = self.doc.add_table(rows=1, cols=2)
        res_table.autofit = False
        res_table.columns[0].width = Inches(3.25)
        res_table.columns[1].width = Inches(3.25)
        r0 = res_table.rows[0]
        self._set_cell_background(r0.cells[0], self.COLORS['patient_info_bg'])
        self._set_cell_background(r0.cells[1], self.COLORS['patient_info_bg'])
        p1 = r0.cells[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.add_run("Screen result: ").bold = True
        res = p1.add_run(overall_risk)
        res.font.bold = True
        res.font.color.rgb = RGBColor.from_string(self.COLORS['low_risk'] if overall_risk == "Low risk" else self.COLORS['high_risk'])
        
        p2 = r0.cells[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(f"Fetal fraction: {ff:.2f}%").bold = True
        
        self._add_summary_table(z_scores)
        
        # Page 2
        self.doc.add_page_break()
        if with_logo: self._add_logo()
        self._add_header_text(2)
        self.doc.add_paragraph() # Spacer
        self._add_detailed_table(z_scores)
        self.doc.add_paragraph().add_run("*As per PCPNDT act, the reference Z scores for sex chromosomal aneuploidies will not be provided.").bold = True
        
        # Page 3
        self.doc.add_page_break()
        if with_logo: self._add_logo()
        self._add_header_text(3)
        self._add_section_header("Interpretation and follow-up")
        for line in ["The results show low risk for all the autosomes and sex chromosomes.", "The fetal fraction was sufficient for analysis.", "A low risk result means that the baby has a low probability of having the condition listed in the report. The results should be communicated by the clinician with appropriate counseling.", "A low-risk result does not eliminate the possibility of aneuploidy or other genetic disorders.", "All results must be correlated with clinical findings, ultrasound, and maternal risk factors."]:
            p = self.doc.add_paragraph(f"• {line}")
            p.style.font.name = 'Segoe UI'
            
        p = self.doc.add_paragraph()
        run = p.add_run("Diagnostic confirmation is strongly recommended if any of the below are:")
        p.style.font.name = 'Segoe UI'
        for line in ["Positive/high-risk NIPT", "Ultrasound abnormalities", "High-risk maternal age or prior affected pregnancy", "No-call/low fetal fraction samples"]:
            p = self.doc.add_paragraph(f"        {line}")
            p.style.font.name = 'Segoe UI'
        
        self._add_section_header("Recommendation")
        p1 = self.doc.add_paragraph("• Genetic counseling is recommended.")
        p1.style.font.name = 'Segoe UI'
        p2 = self.doc.add_paragraph("• If ultrasound/NT scan shows any abnormality or serum markers indicate high risk, invasive testing by amniocentesis with chromosomal microarray is mandated even if NIPS is low risk.")
        p2.style.font.name = 'Segoe UI'
        
        self._add_section_header("Test Method")
        self.doc.add_paragraph(self.METHODOLOGY_TEXT)
        
        self._add_section_header("Test Limitations")
        for p_text in self.LIMITATIONS_TEXT[:2]:
            if "<b>" in p_text:
                p = self.doc.add_paragraph()
                parts = p_text.split("<b>")
                p.add_run(parts[0])
                sub_parts = parts[1].split("</b>")
                p.add_run(sub_parts[0]).bold = True
                p.add_run(sub_parts[1])
            else:
                self.doc.add_paragraph(p_text)
        
        # Page 4
        self.doc.add_page_break()
        if with_logo: self._add_logo()
        self._add_header_text(4)
        
        for p_text in self.LIMITATIONS_TEXT[2:]:
            if "<b>" in p_text:
                p = self.doc.add_paragraph()
                parts = p_text.split("<b>")
                p.add_run(parts[0])
                sub_parts = parts[1].split("</b>")
                p.add_run(sub_parts[0]).bold = True
                p.add_run(sub_parts[1])
            else:
                self.doc.add_paragraph(p_text)
        
        self._add_section_header("Regulatory & Validation Notes")
        self.doc.add_paragraph("Test performance is based on laboratory-validated parameters; results may vary across populations and platforms. Not intended for use in multi-fetal reduction, mosaic embryo transfer, or post-transplant pregnancies.")
        self._add_section_header("Disclaimer")
        self.doc.add_paragraph("No irreversible actions should be taken based solely upon the results of this screening test. The manner in which this information is used to guide patient care is the responsibility of the health care provider, including advising for the need for genetic counseling or diagnostic testing. Any test should be interpreted in the context of all available clinical findings. As with any screening test, any positive result should be confirmed with diagnostic testing.")
        
        # Page 5
        self.doc.add_page_break()
        if with_logo: self._add_logo()
        self._add_header_text(5)
        
        self._add_section_header("Test specifications")
        
        p1 = self.doc.add_paragraph()
        p1.add_run("Table 1: Performance of Sage 32 NIPT workflow in singleton pregnancies").bold = True
        self._add_metrics_table_1()
        
        self.doc.add_paragraph()
        p1_cap = self.doc.add_paragraph("PPV – Positive predictive value; NPV – Negative predictive value; FPR – False positive rate; FNR – False negative rate")
        p1_cap.style.font.name = 'Segoe UI'
        p1_cap.style.font.size = Pt(7)
        p1_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        p2 = self.doc.add_paragraph()
        p2.add_run("Table 2: Performance of Sage 32 NIPT workflow in twin pregnancies").bold = True
        self._add_metrics_table_2()
        
        self.doc.add_paragraph()
        p2_cap = self.doc.add_paragraph("PPV – Positive predictive value; NPV – Negative predictive value; FPR – False positive rate; FNR – False negative rate")
        p2_cap.style.font.name = 'Segoe UI'
        p2_cap.style.font.size = Pt(7)
        p2_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page 6
        self.doc.add_page_break()
        if with_logo: self._add_logo()
        self._add_header_text(6)
        self._add_section_header("Prenatal Testing Algorithm")
        try:
            from nipt_assets import ALGO_CHART_B64
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                tmp_file.write(base64.b64decode(ALGO_CHART_B64))
                algo_path = tmp_file.name
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(algo_path, width=Inches(6.5), height=Inches(3.0))
            os.unlink(algo_path)
        except:
            self.doc.add_paragraph("[Algorithm Flowchart Image Missing]")
        
        self._add_section_header("References")
        refs = [
            ('Bianchi DW, Platt LD, Goldberg JD, Abuhamad AZ, Sehnert AJ, Rava RP. Genome-wide fetal aneuploidy detection by maternal plasma DNA sequencing. ', 'Obstetrics & Gynecology', '. 2012 May 1;119(5):890-901.'),
            ('Chiu RW, Akolekar R, Zheng YW, Leung TY, Sun H, Chan KA, Lun FM, Go AT, Lau ET, To WW, Leung WC. Non-invasive prenatal assessment of trisomy 21 by multiplexed maternal plasma DNA sequencing: large scale validity study. ', 'Bmj', '. 2011 Jan 11;342:c7401.'),
            ('Chiu RW, Lo YM. Noninvasive prenatal diagnosis empowered by high‐throughput sequencing. ', 'Prenatal diagnosis', '. 2012 Apr 1;32(4):401-6.'),
            ('Dugoff L. Cell-Free DNA Screening for Fetal Aneuploidy. ', 'Topics in Obstetrics & Gynecology', '. 2017 Jan 15;37(1):1-7.')
        ]
        for i, (pre, italic, post) in enumerate(refs, 1):
            p = self.doc.add_paragraph(f"{i}. ")
            p.add_run(pre)
            p.add_run(italic).italic = True
            p.add_run(post)
        
        p_rev = self.doc.add_paragraph("This report has been reviewed and approved by:")
        for run in p_rev.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(self.COLORS['blue_header'])
        
        self._add_signature_row()
        
        self.doc.save(self.output_path)
        return self.output_path

    def _set_cell_background(self, cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        tcv = tcPr.find(qn('w:shd'))
        if tcv is not None:
            tcPr.remove(tcv)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

    def _add_section_header(self, text):
        h = self.doc.add_paragraph()
        run = h.add_run(text)
        run.font.name = 'Gill Sans MT Bold'
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor.from_string(self.COLORS['blue_header'])
        
        # Add a bottom border
        pborder = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '48')
        bottom.set(qn('w:color'), self.COLORS['blue_header'])
        pborder.append(bottom)
        h._p.get_or_add_pPr().append(pborder)

    def _add_patient_grid(self, data):
        def fmt_date(d):
            if not d: return ""
            import re
            m = re.match(r'(\d{4})[-\/](\d{1,2})[-\/](\d{1,2})', d)
            if m: return f"{int(m.group(3)):02d}-{int(m.group(2)):02d}-{m.group(1)}"
            m = re.match(r'(\d{1,2})[-\/](\d{1,2})[-\/](\d{4})', d)
            if m: return f"{int(m.group(1)):02d}-{int(m.group(2)):02d}-{m.group(3)}"
            return d
            
        table = self.doc.add_table(rows=6, cols=2)
        mapping = [
            ("Patient name        : ", data.get('name',''), "Specimen                  : ", data.get('specimen','Peripheral blood').title()),
            ("Date of Birth       : ", fmt_date(data.get('dob','')), "PIN                       : ", data.get('pin','')),
            ("Gestational Age     : ", data.get('ga',''), "Sample Number             : ", data.get('sample_id','')),
            ("Pregnancy Type;\nStatus              : ", f"{data.get('preg_type','').title()}; {data.get('preg_status','').title()}".strip('; '), "Sample collection date  : ", fmt_date(data.get('collection_date',''))),
            ("Referring Clinician : ", data.get('clinician',''), "Sample received date  : ", fmt_date(data.get('received_date',''))),
            ("Hospital/Clinic     : ", data.get('hospital',''), "Report date               : ", fmt_date(data.get('report_date', datetime.now().strftime('%d-%m-%Y'))))
        ]
        
        for i, (l1, l2, r1, r2) in enumerate(mapping):
            p1 = table.rows[i].cells[0].paragraphs[0]
            p1.add_run(l1).bold = True
            p1.add_run(l2).bold = True
            p2 = table.rows[i].cells[1].paragraphs[0]
            p2.add_run(r1).bold = True
            p2.add_run(r2).bold = True
            self._set_cell_background(table.rows[i].cells[0], self.COLORS['patient_info_bg'])
            self._set_cell_background(table.rows[i].cells[1], self.COLORS['patient_info_bg'])
            
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.style.font.name = 'Segoe UI'
                    p.style.font.size = Pt(10)

    def _add_summary_table(self, z_scores):
        table = self.doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        hdr0, hdr1 = table.rows[0].cells
        hdr0.text = "Aneuploidies"; hdr1.text = "Results"
        self._set_cell_background(hdr0, self.COLORS['results_header_bg'])
        self._set_cell_background(hdr1, self.COLORS['results_header_bg'])
        for cell in table.rows[0].cells:
            for p in cell.paragraphs:
                for run in p.runs: run.font.bold = True
                p.style.font.name = 'Segoe UI'
        
        table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        targets = [
            ("Chromosome 21", z_scores.get('chr21', 0), 2.8),
            ("Chromosome 18", z_scores.get('chr18', 0), 2.8),
            ("Chromosome 13", z_scores.get('chr13', 0), 2.8),
            ("Sex Chromosomes*", z_scores.get('chrX', 0), 6.0),
            ("Other Autosomes", 0, 6.0)
        ]
        for i, (name, val, thresh) in enumerate(targets, 1):
            row = table.rows[i]
            row.cells[0].text = name
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            self._set_cell_background(row.cells[0], self.COLORS['grey_bg'])
            self._set_cell_background(row.cells[1], self.COLORS['grey_bg'])
            risk = "Low risk" if abs(val) < thresh else "High risk"
            p = row.cells[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_run = p.add_run(risk)
            r_run.font.bold = True
            r_run.font.color.rgb = RGBColor.from_string(self.COLORS['low_risk'] if risk == "Low risk" else self.COLORS['high_risk'])

    def _add_detailed_table(self, z_scores):
        table = self.doc.add_table(rows=23, cols=4)
        table.style = 'Table Grid'
        hdr = ["Chromosome tested", "Z Score", "Test result", "Reference interval"]
        for i, h in enumerate(hdr):
            table.rows[0].cells[i].text = h
            self._set_cell_background(table.rows[0].cells[i], self.COLORS['results_header_bg'])
            p = table.rows[0].cells[i].paragraphs[0]
            run = p.runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(self.COLORS['blue_header'])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i in range(1, 23):
            val = z_scores.get(f'chr{i}', 0)
            thresh = 2.8 if i in [13, 18, 21] else 6.0
            risk = "Low risk" if abs(val) < thresh else "High risk"
            ref = "-6<Z score<2.8" if i in [13, 18, 21] else "-6<Z score<6"
            row = table.rows[i]
            row.cells[0].text = f"  Chromosome {i}"
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            row.cells[1].text = f"{val:.2f}"
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for c in range(4):
                self._set_cell_background(row.cells[c], self.COLORS['grey_bg'])
                row.cells[c].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = row.cells[2].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_run = p.add_run(risk)
            row.cells[3].text = ref
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_signature_row(self):
        try:
            from nipt_assets import SACHIN_SIGN_B64, DIRECTOR_SIGN_B64
            table = self.doc.add_table(rows=3, cols=2)
            for i, b64_str in enumerate([SACHIN_SIGN_B64, DIRECTOR_SIGN_B64]):
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                    tmp_file.write(base64.b64decode(b64_str))
                    sig_path = tmp_file.name
                
                p = table.rows[0].cells[i].paragraphs[0]
                run = p.add_run()
                run.add_picture(sig_path, width=Inches(1.5))
                os.unlink(sig_path)
            
            for i, sig in enumerate(self.SIGNATURES):
                table.rows[1].cells[i].text = sig['name']
                table.rows[2].cells[i].text = sig['title']
                for cell in [table.rows[1].cells[i], table.rows[2].cells[i]]:
                    for p in cell.paragraphs:
                        p.style.font.name = 'Segoe UI'
            for cell in table.rows[1].cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for cell in table.rows[2].cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            table = self.doc.add_table(rows=2, cols=2)
            for i, sig in enumerate(self.SIGNATURES):
                table.rows[0].cells[i].text = sig['name']
                table.rows[1].cells[i].text = sig['title']
                for cell in [table.rows[0].cells[i], table.rows[1].cells[i]]:
                    for p in cell.paragraphs:
                        p.style.font.name = 'Segoe UI'
            for cell in table.rows[0].cells:
                for p in cell.paragraphs:
                    for run in p.runs: run.font.bold = True

    def _add_metrics_table_1(self):
        table = self.doc.add_table(rows=6, cols=7)
        table.style = 'Table Grid'
        hdr = ["Condition", "Sensitivity", "Specificity", "PPV", "NPV", "FPR", "FNR"]
        for i, h in enumerate(hdr):
            table.rows[0].cells[i].text = h
            self._set_cell_background(table.rows[0].cells[i], self.COLORS['results_header_bg'])
            for r_ in table.rows[0].cells[i].paragraphs[0].runs: r_.font.bold = True
            table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.rows[0].cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        try:
            from docx.shared import Inches
            table.rows[0].height = Inches(0.4)
        except:
            pass
        data = [
            ["Trisomy 21", "99%\n95%CI:99.49-100%", ">99%\n95%CI:99.99-100%", ">99.99%", "99.99%", "<0.1%", "0.1%"],
            ["Trisomy 18", "99%\n95%CI:96.82-99.99%", ">99%\n95%CI:99.99-100%", ">99.99%", "99.99%", "<0.1%", "0.58%"],
            ["Trisomy 13", ">99%\n95%CI:96.87-100%", ">99%\n95%CI:99.99-100%", ">99.99%", ">99.99%", "<0.1%", "<0.1%"],
            ["Sex chromosomal\naneuploidies", ">98 %\n95%CI:99.31-100%", ">98%\n95%CI:99.99-100%", "99.81%", ">99.99%", "0.0014%", "<0.01%"],
            ["Other autosomal\naneuploidies", ">90%\n95%CI:98.03-100%", ">90%\n95%CI:99.99-100%", ">99.99%", ">99.99%", "<0.01%", "<0.01%"]
        ]
        for i, r in enumerate(data, 1):
            self._set_cell_background(table.rows[i].cells[0], 'E2E2E2')
            for c in range(1, 7):
                self._set_cell_background(table.rows[i].cells[c], self.COLORS['grey_bg'])
            for c, txt in enumerate(r):
                if '\n' in txt:
                    parts = txt.split('\n')
                    table.rows[i].cells[c].text = parts[0] + " "
                    p = table.rows[i].cells[c].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.style.font.size = Pt(8)
                    run2 = p.add_run(parts[1])
                    run2.font.size = Pt(5.5)
                    table.rows[i].cells[c].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                else:
                    table.rows[i].cells[c].text = txt
                    table.rows[i].cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.rows[i].cells[c].paragraphs[0].style.font.size = Pt(8)
                    table.rows[i].cells[c].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _add_metrics_table_2(self):
        table = self.doc.add_table(rows=6, cols=7)
        table.style = 'Table Grid'
        hdr = ["Condition", "Sensitivity", "Specificity", "PPV", "NPV", "FPR", "FNR"]
        for i, h in enumerate(hdr):
            table.rows[0].cells[i].text = h
            self._set_cell_background(table.rows[0].cells[i], self.COLORS['results_header_bg'])
            for r_ in table.rows[0].cells[i].paragraphs[0].runs: r_.font.bold = True
            table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.rows[0].cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        try:
            from docx.shared import Inches
            table.rows[0].height = Inches(0.4)
        except:
            pass
        data = [
            ["Trisomy 21", ">98%", ">98%", ">99.99%", ">99.99%", "<0.01%", "<0.01%"],
            ["Trisomy 18", ">98%", ">98%", ">98%", ">99.99%", "<0.01%", "<0.01%"],
            ["Trisomy 13", ">98%", ">98%", ">98%", ">99.99%", "<0.01%", "<0.01%"],
            ["Sex chromosomal\naneuploidies", ">97%", ">97%", ">97%", ">99.99%", "0.0014%", "<0.01%"],
            ["Other autosomal\naneuploidies", ">95%", ">95%", ">95%", ">99.99%", "<0.01%", "<0.01%"]
        ]
        for i, r in enumerate(data, 1):
            self._set_cell_background(table.rows[i].cells[0], 'E2E2E2')
            for c in range(1, 7):
                self._set_cell_background(table.rows[i].cells[c], self.COLORS['grey_bg'])
            for c, txt in enumerate(r):
                table.rows[i].cells[c].text = txt
                table.rows[i].cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.rows[i].cells[c].paragraphs[0].style.font.size = Pt(8)
                table.rows[i].cells[c].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

if __name__ == "__main__":
    g = NIPTDocxGenerator("test_nipt_v3.docx")
    g.generate({'name': 'TEST'}, {f'chr{i}': 0.1 for i in range(1, 23)})
    print("V3 DOCX Corrected.")
